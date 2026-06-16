import os
import json
import pandas as pd
from docx import Document

# Define directory structure
RAW_DATA_DIR = os.path.join("data", "raw")
os.makedirs(RAW_DATA_DIR, exist_ok=True)

def save_metadata(file_path, title, tags, doc_type):
    """Generates a sidecar JSON metadata file for tracking."""
    metadata = {
        "title": title,
        "tags": tags,
        "type": doc_type,
        "date": "2026-06-09"
    }
    meta_path = file_path + ".json"
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=4)
    print(f"Generated metadata for: {os.path.basename(file_path)}")

def generate_markdown():
    """Generates Markdown files for company overview and policies."""
    content = """# FloCard Engineering Core Overview

Welcome to the FloCard Engineering ecosystem. This document serves as the primary technical onboarding guide for incoming developers.

## Our Mission
Our core objective is to redefine digital identity management and contact synchronization pipelines. We build scalable, highly secure systems that allow seamless P2P profile exchanges.

## Sprint Operations
- Sprints operate on strict 15-day development cycles.
- Daily standups are hosted on Microsoft Teams channels promptly at 9:30 AM IST.
- Jira tickets must be accurately categorized (Epic, Story, Bug, Task).
"""
    file_path = os.path.join(RAW_DATA_DIR, "engineering_overview.md")
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content.strip())
    save_metadata(file_path, "FloCard Engineering Overview", ["onboarding", "engineering", "sprints"], "Markdown")

def generate_docx():
    """Generates a DOCX file for standard operating procedures (SOP)."""
    doc = Document()
    doc.add_heading("SOP: Code Review and Git Workflow", level=1)
    
    doc.add_heading("1. Branching Strategy", level=2)
    p1 = doc.add_paragraph()
    p1.add_run("All engineering work must originate from isolated short-lived branches. Use the convention ")
    p1.add_run("feature/your-feature-name").bold = True
    p1.add_run(". Do not push direct commits onto the dev or main branch pathways.")
    
    doc.add_heading("2. Pull Request Guidelines", level=2)
    doc.add_paragraph("Before requesting a merge from feature/* into dev, ensure that all unit tests execute and pass successfully locally. A minimum of two peer code reviews is required to clear the quality gate before ultimate branch merging.")
    
    file_path = os.path.join(RAW_DATA_DIR, "git_workflow_sop.docx")
    doc.save(file_path)
    save_metadata(file_path, "Git Workflow SOP", ["git", "sop", "code-review"], "DOCX")

def generate_csv():
    """Generates a CSV file containing FAQ metrics."""
    faq_data = {
        "Question": [
            "What is the daily cutoff time for logging progress updates in Jira?",
            "How do I request help if my ticket is facing an unexpected blocker?",
            "Where should testing team metrics and FOU certificates be stored?"
        ],
        "Answer": [
            "All Jira tickets must be updated completely by 3:20 PM IST daily before office closure.",
            "Tag team members in your Jira ticket comments and post the specific block alerts onto your specific Microsoft Teams channel.",
            "All completed software test cases, validation logs, and Fitness of Use (FOU) certifications are kept permanently inside DevOps archives."
        ]
    }
    df = pd.DataFrame(faq_data)
    file_path = os.path.join(RAW_DATA_DIR, "operational_faqs.csv")
    df.to_csv(file_path, index=False, encoding="utf-8")
    save_metadata(file_path, "Operational FAQs Log", ["jira", "faq", "operations"], "CSV")

if __name__ == "__main__":
    print("🚀 Starting synthetic corporate data generation...")
    generate_markdown()
    generate_docx()
    generate_csv()
    print("✅ Synthetic data and tracking metadata catalogs successfully built inside data/raw/!")